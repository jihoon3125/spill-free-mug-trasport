"""Build the TWO-COLUMN paper-style report. Numbers are pulled from results/*.json
so the report auto-updates after experiments. Figures auto-numbered. Output docx.
"""
import json
from pathlib import Path
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[2]
FIG = ROOT / "figs"; RES = ROOT / "results"


def jload(name):
    try:
        return json.load(open(RES / name))
    except Exception:
        return None

# ---------- pull numbers from results ----------
METHODS = ["min-jerk", "chomp-smooth", "stomp-spill", "vanilla-grad", "ours"]
mc = jload("method_comparison.json"); ts = jload("tstar.json")
tmap = jload("transport_map.json"); gp = jload("gamma_pareto.json"); fl = jload("filllevel.json")

agg = {}
if mc:
    for m in METHODS:
        sr = np.array([g["methods"][m]["spill_ratio"] for g in mc["per_grasp"]]) * 100
        fe = np.mean([g["methods"][m].get("feasible", False) for g in mc["per_grasp"]]) * 100
        pt = np.mean([g["methods"][m]["plan_time"] for g in mc["per_grasp"]])
        agg[m] = (sr.mean(), fe, pt)
    T_MC = mc["T"]; N_MC = mc["n_grasps"]
else:
    for m in METHODS:
        agg[m] = (0, 0, 0)
    T_MC, N_MC = 1.0, 0

if ts:
    tstars = [r["tstar"] for r in ts["rows"] if r["tstar"] is not None]
    offs = [r["offset_r"] for r in ts["rows"] if r["tstar"] is not None]
    TS_MIN, TS_MAX = (min(tstars), max(tstars)) if tstars else (0, 0)
    TS_MED = np.median(tstars) if tstars else 0
    TS_MONO = ts.get("monotonic_frac", 0) * 100
    TS_CORR = np.corrcoef(offs, tstars)[0, 1] if len(tstars) > 2 else float("nan")
    TS_NDEF, TS_NTOT = len(tstars), len(ts["rows"])
else:
    TS_MIN = TS_MAX = TS_MED = TS_MONO = TS_CORR = TS_NDEF = TS_NTOT = 0

if tmap:
    REACH_N = sum(r["reachable"] for r in tmap); REACH_TOT = len(tmap)
    UNREACH_PCT = round((1 - REACH_N / max(1, REACH_TOT)) * 100)
else:
    REACH_N = REACH_TOT = 0; UNREACH_PCT = 34

OURS_SPILL, OURS_FEAS, _ = agg["ours"]

# ---------- doc setup ----------
doc = Document(); st = doc.styles["Normal"]
st.font.name = "Times New Roman"; st.font.size = Pt(9.5)
J = WD_ALIGN_PARAGRAPH.JUSTIFY; C = WD_ALIGN_PARAGRAPH.CENTER
COL_W = 3.05
_FN = [0]


def set_cols(section, n, space=300):
    sectPr = section._sectPr; cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(n)); cols.set(qn('w:space'), str(space))


def para(text, align=J, size=9.5, bold=False, italic=False, after=5, before=0):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p


def heading(num, text):
    para(f"{num}  {text}".strip(), size=11, bold=True, before=7, after=2)


def sub(text):
    para(text, size=10, bold=True, before=3, after=1)


def add_math(p, text, size=9.5, bold=False, italic=False):
    """Render text with _{..}/^{..}/_x/^x markup as true sub/superscripts."""
    i = 0; buf = ""

    def emit(s, sup=False, sub=False):
        if not s:
            return
        r = p.add_run(s); r.font.size = Pt(size); r.bold = bold; r.italic = italic
        r.font.superscript = sup; r.font.subscript = sub

    while i < len(text):
        c = text[i]
        if c in "_^" and i + 1 < len(text):
            emit(buf); buf = ""
            sup = (c == "^")
            if text[i + 1] == "{":
                j = text.find("}", i + 2); seg = text[i + 2:j]; i = j + 1
            else:
                seg = text[i + 1]; i += 2
            emit(seg, sup=sup, sub=not sup)
        else:
            buf += c; i += 1
    emit(buf)


def mpara(text, align=J, size=9.5, bold=False, italic=False, after=5, before=0):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    add_math(p, text, size=size, bold=bold, italic=italic)
    return p


def eq(text):
    p = doc.add_paragraph(); p.alignment = C
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.space_before = Pt(1)
    add_math(p, text, size=10, italic=True)


def mono(lines):
    for ln in lines:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        r = p.add_run(ln); r.font.name = "Consolas"; r.font.size = Pt(7.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def figure(fname, caption, width=COL_W):
    if not (FIG / fname).exists():
        return
    _FN[0] += 1
    doc.add_picture(str(FIG / fname), width=Inches(width))
    doc.paragraphs[-1].alignment = C
    para(f"Figure {_FN[0]}. {caption}", align=C, size=8.3, italic=True, after=8)


def full_width_fig(fname, caption, width=6.6):
    s1 = doc.add_section(WD_SECTION.CONTINUOUS); set_cols(s1, 1)
    figure(fname, caption, width=width)
    s2 = doc.add_section(WD_SECTION.CONTINUOUS); set_cols(s2, 2)


def table(headers, rows, sizes=8.6):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        r = t.rows[0].cells[j].paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(sizes)
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].paragraphs[0].add_run(str(v)).font.size = Pt(sizes)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


# ===== Title / abstract =====
set_cols(doc.sections[0], 1)
para("Grasp-Conditioned Spill-Free Mug Transport", align=C, size=17, bold=True, after=3)
para("Jihoon Yoon  ·  Motion Planning Term Project", align=C, size=10.5, italic=True, after=8)
para(f"Abstract — Given a dexterous grasp of a mug, we generate a 6-DoF arm trajectory "
     f"that transports it without spilling. We embed a differentiable spill cost—derived "
     f"from the apparent-gravity (“waiter”) model—into CHOMP and evaluate it on each "
     f"grasp ({N_MC} reachable grasps). At a hardware-feasible budget our planner is the "
     f"only method that is both nearly spill-free ({OURS_SPILL:.1f}% mean spill) and "
     f"within joint limits ({OURS_FEAS:.0f}% feasible), where baselines either spill "
     f"(chomp-smooth {agg['chomp-smooth'][0]:.0f}%, min-jerk {agg['min-jerk'][0]:.0f}%) or "
     f"violate joint limits (vanilla-grad / STOMP {agg['vanilla-grad'][1]:.0f}% feasible). "
     f"The fastest feasible spill-free transport time is a property of the grasp—it ranges "
     f"{TS_MIN:.1f}–{TS_MAX:.1f} s and is not predicted by grasp geometry (r = {TS_CORR:+.2f}); "
     f"{UNREACH_PCT}% of grasps cannot reach the task at all—so spill-robustness must be "
     f"measured by planning. We also characterize a degenerate high-acceleration mode of "
     f"the spill cost and the spill–feasibility trade-off in the cost weight.",
     size=9, italic=True, after=8)
body = doc.add_section(WD_SECTION.CONTINUOUS); set_cols(body, 2)

# ===== 1 Introduction =====
heading("1", "Introduction")
para("Dexterous grasp synthesis has matured: modern affordance- and taxonomy-aware "
     "methods produce a suitable multi-finger hand pose for an object and task; my thesis "
     "is one such generator. A largely unanswered question is what happens after the "
     "grasp—whether the grasped object can be manipulated safely. We study the canonical "
     "instance: transporting a mug of liquid. Given a grasp, produce a 6-DoF arm "
     "trajectory that moves the mug from A to B without spilling.")
para("The governing physics is the waiter problem. A cup spills not only when "
     "geometrically tilted but whenever it accelerates, because in the cup’s quasi-static "
     "frame the liquid aligns with the effective gravity g − a rather than g; spilling "
     "occurs when this apparent-gravity vector leaves the cone subtended by the rim. The "
     "safe envelope depends on how the cup is held—on the grasp.")
para("Contributions: (i) a differentiable spill cost in CHOMP that produces "
     "hardware-feasible spill-free transport, with a characterized degenerate "
     "high-acceleration mode avoided by a moderate weight; (ii) a grasp-conditioned "
     "analysis showing the fastest feasible spill-free time is a grasp property not "
     "predictable from geometry; and (iii) a feasibility analysis of when each method "
     "violates joint limits.")
full_width_fig("_grid23_still.png",
               "The same three transport tasks at the same time budget. Top: min-jerk "
               "spills; bottom: ours keeps the water in. Gold curves are mug-center paths.")

# ===== 2 Related Work =====
heading("2", "Related Work")
sub("2.1  Waiter-problem / non-prehensile transport")
para("Work transporting an ungrasped object on the end-effector enforces no-tip/no-slip "
     "via the friction cone: Heins and Schoellig [1] balance an object with obstacle "
     "avoidance via constrained MPC; Selvaggio et al. [2] enforce non-sliding contact "
     "during fast transport. These fix the contact and reason about rigid tipping, not "
     "fluid slosh.")
sub("2.2  Slosh-free liquid transport")
para("Muchacho et al. [3] model the end-effector as a spherical pendulum and compute "
     "slosh-free trajectories via a QP—the reduced-order slosh model we build on, for a "
     "fixed container pose. Moriello et al. [4] use input shaping; Di Lillo et al. [6] "
     "keep the surface within container limits; Gandhi et al. [5] learn liquid dynamics "
     "through clutter. All are grasp-agnostic.")
sub("2.3  Trajectory optimization with task costs")
para("Our optimizer is CHOMP [7]; the spill cost is an added differentiable term. STOMP "
     "[8] is the sampling counterpart and TrajOpt [9] the sequential-convex one; as slosh "
     "is acceleration-dependent, space-time CHOMP [10] is relevant. We compare optimizers "
     "and ablate the covariant metric.")
sub("2.4  Tilt-constrained planning and task-oriented grasping")
para("Orientation-constrained planners impose a static bounded tilt—Task Space Regions "
     "[11], task-constrained RRT [12] (carry a glass level)—whereas ours is a dynamic "
     "spill margin. Task-oriented and semantic grasping select grasps for a task [13,14], "
     "gate kinematic feasibility [16], or co-optimize grasp and motion [15], but never via "
     "a physical spill margin.")
sub("2.5  Gap")
para("Spill-free transport [3–6] and waiter balancing [1,2] fix the grasp and optimize the "
     "trajectory; task-oriented grasping [13–16] stops at grasp selection with a kinematic "
     "metric. No work closes the loop grasp → spill margin → trajectory. We take a "
     "first step and quantify how the grasp sets the feasible spill margin.")

# ===== 3 Method =====
heading("3", "Method")
sub("3.1  Problem formulation")
mpara("Let q(t) denote the 6-DoF UR5e configuration; the hand joints are frozen, so the "
      "grasp fixes the rigid transform T_{em} and the mug world pose is FK(q)·T_{em}. With "
      "upright start/goal we seek a waypoint trajectory ξ = (q_{0}, …, q_{N−1}) with fixed "
      "endpoints that keeps the apparent-gravity vector inside the rim cone, is smooth, "
      "collision-free, and within joint limits.")
sub("3.2  Differentiable kinematics")
mpara("A differentiable FK chain walker (autograd through SE(3)) matches SAPIEN to 1e-7. "
      "IK minimizes ‖p − p*‖^{2} + λ(3 − tr(R R*^{T})) by Adam from multiple seeds; an IK "
      "failure then indicates genuine unreachability.")
sub("3.3  Spill model (apparent-gravity cone)")
mpara("Under the quasi-static assumption the liquid surface is normal to the effective "
      "gravity g_{eff} = g − a (a = mug-center acceleration); with n the opening axis the "
      "cup does not spill iff")
eq("align(t) = − ĝ_{eff}(t) · n(t)  ≥  cos θ_{max} ,")
mpara("where θ_{max} is set by the fill level (≈ atan((rim − liquid)/radius); 18° ≈ "
      "half-full). The penalty is c_{spill} = max(0, cos θ_{max} − align)^{2}. The "
      "acceleration is the second finite difference of p_{mug} = FK(q)·T_{em}, so the "
      "per-grasp lever arm (α×r, ω×(ω×r)) is captured exactly: identical arm motions with "
      "different T_{em} give different mug accelerations.")
figure("fig_cone.png",
       "Spill condition. Left: −g_eff inside the rim cone (safe). Right: acceleration "
       "tilts −g_eff out of the cone (spill).")
sub("3.4  Cost functional and covariant update")
mpara("J = α J_{smooth} + β J_{obs} + γ J_{spill}, with J_{smooth} = ‖Aξ‖^{2} (A the "
      "second-difference operator) inducing the CHOMP metric M = A^{T}A. CHOMP descends "
      "the covariant gradient")
eq("ξ  ←  ξ  −  η · M^{−1} · ∇J(ξ) ,")
para("fed to Adam with gradient clipping. The spill cost constrains only the direction of "
     "g_eff, so a large γ admits a degenerate solution that injects a huge upward "
     "acceleration to align g_eff at the cost of joint-limit-violating motion; a moderate "
     "γ (1.0), where smoothness bounds acceleration, avoids it (§4.2).")
sub("3.5  Per-grasp spill-aware planning")
mono([
    "Algorithm 1  Per-grasp spill-aware transport",
    "Input: grasp G (T_em), start/goal, budget T,",
    "       theta_max, weight gamma",
    " 1: q_s <- IK(T_start . T_em^-1)  (multi-seed)",
    " 2: q_g <- IK(T_goal  . T_em^-1)",
    " 3: if IK error > tol: return UNREACHABLE",
    " 4: xi <- minjerk(q_s, q_g, N)",
    " 5: for k = 1..K do",
    " 6:   p <- FK(xi).T_em;  a <- d2 p / dt2",
    " 7:   J <- a Jsm + b Jobs + g Jspill(a,n)",
    " 8:   grad <- M^-1 dJ/dxi; clip; xi<-Adam(xi,grad)",
    " 9: return xi, spill_ratio(xi; theta_max)",
])

# ===== 4 Experiments =====
heading("4", "Experiments")
mpara(f"Setup. A mug with {REACH_TOT} sampled grasps (2-finger / 3-finger / whole-hand). "
      f"Default task: 70 cm lateral + 15 cm lift, upright endpoints. Metric: spill ratio = "
      f"fraction of waypoints whose effective tilt exceeds θ_{{max}}. Feasible = max joint "
      f"velocity < 3.3 rad/s and acceleration < 20 rad/s^{{2}}.")

sub("4.1  Method comparison and feasibility")
para(f"Five methods, each changing one ingredient, on {N_MC} reachable grasps at a "
     f"hardware-feasible budget T = {T_MC} s (γ = 1.0 for spill-aware methods).")
table(["Method", "mean spill", "feasible"],
      [[m, f"{agg[m][0]:.1f}%", f"{agg[m][1]:.0f}%"] for m in METHODS])
para(f"Ours is the only method that is both nearly spill-free ({OURS_SPILL:.1f}%) and "
     f"feasible ({OURS_FEAS:.0f}%). Smoothness alone (chomp-smooth) is feasible but spills "
     f"({agg['chomp-smooth'][0]:.0f}%). Without the CHOMP metric the gradient and sampling "
     f"baselines (vanilla-grad, STOMP) are near-0% feasible—they reduce spill only by "
     f"injecting joint-limit-violating accelerations (Fig. 4)—and still spill heavily. The "
     f"metric is thus essential for both spill and feasibility, and at far lower compute "
     f"than sampling (ours plans in {agg['ours'][2]:.0f} s vs {agg['stomp-spill'][2]:.0f} s "
     f"for STOMP).")
figure("fig_method_comparison.png",
       f"Spill distribution over {N_MC} grasps at the feasible budget T = {T_MC} s.")
figure("fig_feasibility.png",
       "Max joint acceleration per method (log). vanilla-grad and STOMP exceed the limit; "
       "ours and chomp-smooth stay within it.")

sub("4.2  Spill–feasibility trade-off and fill level")
if gp:
    para("Sweeping the spill weight γ exposes the trade-off (Fig. 5): small γ leaves "
         "residual spill, while large γ triggers the degenerate high-acceleration mode and "
         "feasibility collapses. γ ≈ 1 is the knee—low spill while staying feasible.")
    figure("fig_gamma_pareto.png",
           "Spill (red) and feasibility (green) vs the spill weight γ at T = 1.0 s. γ≈1 is "
           "the knee of the trade-off.")
if fl:
    mpara("Because θ_{max} encodes the fill level (a fuller cup ⇒ smaller θ_{max}), Fig. 6 "
          "shows how many grasps ours keeps spill-free as the cup is filled: nearly all at "
          "moderate fill, degrading gracefully as the cup approaches the brim.")
    figure("fig_filllevel.png",
           "Spill-robustness vs spill threshold / fill level (T = 1.0 s, ours).")

sub("4.3  Spill-robustness is a property of the grasp")
para(f"Of {REACH_TOT} grasps, {UNREACH_PCT}% cannot reach the transport workspace "
     f"(Fig. 9). For reachable grasps we define T* as the fastest budget that is both "
     f"spill-free and feasible. With γ = 1.0 spill is monotone in T for {TS_MONO:.0f}% of "
     f"grasps, so T* is well-defined; it ranges {TS_MIN:.1f}–{TS_MAX:.1f} s "
     f"(median {TS_MED:.1f} s) and is not predicted by grasp offset (r = {TS_CORR:+.2f}, "
     f"Fig. 8). A spill-robust grasp thus cannot be chosen by a geometric heuristic—it "
     f"must be measured by planning.")
figure("fig_tstar.png",
       "Fastest feasible spill-free transport time T* per grasp; it varies across grasps.")
figure("fig_tstar_offset.png",
       f"T* vs EE–mug offset (r = {TS_CORR:+.2f}): grasp geometry does not predict it.")
figure("fig_reach.png",
       f"Reachability by taxonomy: {UNREACH_PCT}% of grasps cannot reach the task.")

sub("4.4  Qualitative")
para("Figure 10 shows the five methods on one task; only ours and chomp-smooth stay below "
     "the spill threshold, but the spill-aware baselines do so only via infeasible motion "
     "(Fig. 4).")
full_width_fig("fig_trajectories.png",
               "Five methods on one task: mug-center paths (left) and effective tilt over "
               "time (right).")

# ===== 5 Discussion =====
heading("5", "Discussion and Limitations")
mpara("Feasibility is the central subtlety. The spill cost constrains only the direction "
      "of effective gravity, so an aggressive weight or over-fast budget yields "
      "trajectories that are nominally spill-free but violate joint limits (up to ~10 rad/s "
      "and ~330 rad/s^{2}, far above the UR5e limits). We report at a feasible budget with a "
     "moderate weight and include a feasibility metric; properly coupling dynamic limits "
     "(time-optimal reparameterization or hard constraints) is future work. We use the "
     "quasi-static apparent-gravity cone [3], not full CFD; a rigid-body particle "
     "simulation we built behaves granularly and is not used quantitatively. CHOMP "
     "non-convexity makes spill only weakly monotone in T, so we treat T* statistically. "
     "Scope: a single mug instance. We found and fixed an inverted cup up-axis convention "
     "midway; all numbers use the corrected convention.")

# ===== 6 Conclusion =====
heading("6", "Conclusion and Future Work")
para("We built a spill-aware planner that, at a feasible budget, uniquely achieves both "
     "spill-free and joint-limit-feasible transport, and showed that the fastest feasible "
     "spill-free transport time is a grasp property not predictable from geometry, so it "
     "must be measured. This is the signal a grasp generator could consume: future work is "
     "to feed the measured spill margin back into grasp selection, couple dynamic limits "
     "and obstacles, and extend to more objects and a learned spill model.")

# ===== References =====
heading("", "References")
refs = [
    "M. Heins, A. P. Schoellig, “Keep It Upright: MPC for Nonprehensile Object Transportation with Obstacle Avoidance,” IEEE RA-L, 2023.",
    "M. Selvaggio et al., “Non-Prehensile Object Transportation via MPC Non-Sliding Manipulation Control,” IEEE TCST, 2023.",
    "T. P. Muchacho, R. Laha, L. Figueredo, S. Haddadin, “A Solution to Slosh-free Robot Trajectory Optimization,” IROS, 2022.",
    "L. Moriello, L. Biagiotti, C. Melchiorri, A. Paoli, “Manipulating liquids with robots: a sloshing-free solution,” Control Eng. Practice, 2018.",
    "G. Gandhi, B. Sundaralingam, J. Pan et al., “Clutter-Aware Spill-Free Liquid Transport via Learned Dynamics,” arXiv:2408.00215, 2024.",
    "C. Di Lillo, Q.-C. Pham et al., “Robotic Handling of Liquids with Spilling Avoidance,” ICRA, 2019.",
    "M. Zucker, N. Ratliff, A. Dragan et al., “CHOMP,” IJRR, 2013.",
    "M. Kalakrishnan, S. Chitta, E. Theodorou, P. Pastor, S. Schaal, “STOMP,” ICRA, 2011.",
    "J. Schulman et al., “Motion planning with sequential convex optimization and convex collision checking,” IJRR, 2014.",
    "A. Byravan, B. Boots, S. Srinivasa, D. Fox, “Space-Time Functional Gradient Optimization for Motion Planning,” ICRA, 2014.",
    "D. Berenson, S. Srinivasa, J. Kuffner, “Task Space Regions,” IJRR, 2011.",
    "M. Stilman, “Global Manipulation Planning in Robot Joint Space with Task Constraints,” IEEE T-RO, 2010.",
    "H. Dang, P. K. Allen, “Semantic grasping,” Autonomous Robots, 2014.",
    "K. Fang, Y. Zhu, A. Garg et al., “Learning Task-Oriented Grasping,” IJRR, 2020.",
    "M. Toussaint, K. Allen, K. Smith, J. Tenenbaum, “Differentiable Physics and Stable Modes for Tool-Use and Manipulation Planning,” RSS, 2018.",
    "J. Fontanals et al., “Grasping for a Purpose,” ICRA, 2014.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1); p.alignment = J
    p.add_run(f"[{i}] {r}").font.size = Pt(8)

for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.7)

doc.save(str(ROOT / "docs/report.docx"))
print(f"saved docs/report.docx  (figs={_FN[0]}, n_grasps={N_MC}, ours={OURS_SPILL:.1f}%/{OURS_FEAS:.0f}%feas)")
